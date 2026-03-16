import mimetypes
import os

import markdown  # type: ignore
from bs4 import BeautifulSoup
from docx import Document as DocxDocument
from pypdf import PdfReader

from src.core.config import config
from src.core.exceptions import IngestionError, UnsupportedFileTypeError
from src.core.logging import get_logger
from src.core.schemas import IngestedDocument

logger = get_logger(__name__)

# Maximum file size accepted by the ingestion service (default 50 MB).
MAX_FILE_SIZE_BYTES: int = int(os.getenv("MAX_UPLOAD_MB", "50")) * 1024 * 1024


class DocumentIngestionService:
    """Handles raw file reading and textual extraction across multiple formats."""

    def __init__(self) -> None:
        self.supported_extensions: list[str] = config.get_nested(
            "pipeline.ingestion.supported_extensions",
            [".pdf", ".docx", ".txt", ".md", ".html", ".csv", ".json"],
        )

    def ingest_file(self, file_path: str) -> IngestedDocument:
        """Read *file_path* from disk and extract its raw text and basic metadata.

        Args:
            file_path: Absolute path to the file on disk.

        Returns:
            An :class:`IngestedDocument` containing the raw text and metadata.

        Raises:
            :class:`FileNotFoundError`: If the file does not exist.
            :class:`UnsupportedFileTypeError`: If the extension is not supported.
            :class:`IngestionError`: If text extraction fails for any reason.
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        stat = os.stat(file_path)
        if stat.st_size > MAX_FILE_SIZE_BYTES:
            max_mb = MAX_FILE_SIZE_BYTES // (1024 * 1024)
            raise IngestionError(
                f"File '{os.path.basename(file_path)}' exceeds the maximum upload size "
                f"of {max_mb} MB ({stat.st_size:,} bytes received)."
            )

        ext = os.path.splitext(file_path)[1].lower()
        if ext not in self.supported_extensions:
            raise UnsupportedFileTypeError(
                f"Unsupported file type '{ext}'. "
                f"Supported: {', '.join(self.supported_extensions)}"
            )

        raw_text = self._extract_text(file_path, ext)
        if not raw_text:
            raise IngestionError(
                f"Text extraction returned empty content for '{os.path.basename(file_path)}'. "
                "The file may be empty, corrupt, or contain only images."
            )

        logger.info(
            "document_ingested",
            filename=os.path.basename(file_path),
            ext=ext,
            size_bytes=stat.st_size,
            chars=len(raw_text),
        )

        return IngestedDocument(
            source=file_path,
            filename=os.path.basename(file_path),
            type=mimetypes.guess_type(file_path)[0] or "unknown",
            size_bytes=stat.st_size,
            raw_text=raw_text,
            metadata={"extension": ext},
        )

    def _extract_text(self, file_path: str, ext: str) -> str:
        text = ""
        try:
            if ext == ".pdf":
                reader = PdfReader(file_path)
                for page in reader.pages:
                    extracted = page.extract_text()
                    if extracted:
                        text += extracted + "\n"
            elif ext == ".docx":
                doc = DocxDocument(file_path)
                text = "\n".join([para.text for para in doc.paragraphs])
            elif ext == ".md":
                with open(file_path, encoding="utf-8") as f:
                    md_text = f.read()
                    html = markdown.markdown(md_text)
                    text = BeautifulSoup(html, "html.parser").get_text()
            elif ext in [".txt", ".csv", ".json", ".html", ".py", ".js", ".ts", ".java", ".cpp"]:
                with open(file_path, encoding="utf-8", errors="replace") as f:
                    if ext == ".html":
                        text = BeautifulSoup(f.read(), "html.parser").get_text()
                    else:
                        text = f.read()
        except Exception as exc:
            raise IngestionError(
                f"Failed to extract text from '{os.path.basename(file_path)}': {exc}"
            ) from exc
        return text.strip()


ingestion_service = DocumentIngestionService()
